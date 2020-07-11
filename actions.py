# This files contains your custom actions which can be used to run
# custom Python code.
#
# See this guide on how to implement these action:
# https://rasa.com/docs/rasa/core/actions/#custom-actions/


# This is a simple example for a custom action which utters "Hello World!"

from typing import Any, Text, Dict, List, Union

from rasa_sdk import Action, Tracker
from rasa_sdk.executor import CollectingDispatcher
from rasa_sdk.events import SlotSet
from rasa_sdk.forms import FormAction

import mysql.connector

from urllib.request import urlopen
import json

import googlemaps
import time

from docx import Document
from docx.shared import Inches

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders


class ActionExerciceSearch(Action):

    def name(self) -> Text:
        return "action_exercice_search"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        mydb = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="p@ss123",
            database="testdatabase"
            # auth_plugin='mysql_native_password'
        )

        print('action_exercice_search')
        exercice = tracker.get_slot("exercice_type")
        muscularGroup = tracker.get_slot("muscular_group")

        cursor = mydb.cursor()
        query = "SELECT * FROM exercises where exercise_group = '%s'" % muscularGroup
        try:
            cursor.execute(query)
            if cursor.execute(query) == 0:
                print("The data was extracted")
        except:
            print("mysql")

        # dispatcher.utter_message("Here is the muscle group of the {}:{}".format(exercice,muscularGroup))
        dispatcher.utter_message("Here are some exercises for {}:".format(muscularGroup))
        records = cursor.fetchall()

        for counter, exercices in enumerate(records):
            dispatcher.utter_message(str(counter + 1))
            dispatcher.utter_message(str(exercices[0]))
            dispatcher.utter_message(
                "You should do {} reps with {} sets. {}".format(exercices[4], exercices[5], exercices[6]))

        return []


class ActionConfirmUserEmail(Action):

    def name(self) -> Text:
        return "action_confirm_user_email"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

        print("action_confirm_user_email")
        mydb = mysql.connector.connect(host="localhost", user="root", passwd="p@ss123",
                                       database="testdatabase")  # # auth_plugin='mysql_native_password'
        userEmail = tracker.latest_message.get('text')
        print(userEmail)
        query = "select user_name from users where user_email = '{}';".format(userEmail)
        cursor = mydb.cursor()
        print(query)

        try:
            x = cursor.execute(query)
            if x == 0:
                print("Sorry, could not find you in the DB")
                dispatcher.utter_message("Sorry I could not find you by your email! :( ")
            else:
                result = cursor.fetchone()
                
                dispatcher.utter_message("Welcome back, {} !".format(str(result[0])))
                dispatcher.utter_message("How can I help you?")
        except:
            dispatcher.utter_message("Sorry I could not find you by your email! :( ")
        return [SlotSet("user_email", userEmail)]


class ActionFormUserInfo(FormAction):

    def name(self) -> Text:
        return "form_user"

    @staticmethod
    def required_slots(tracker: Tracker) -> List[Text]:
        return ["user_name", "user_age", "user_weight", "user_height", "user_sex", "user_scope", "user_times_at_gym",
                "user_email"]

    def slot_mappings(self) -> Dict[Text, Union[Dict, List[Dict]]]:
        """A dictionary to map required slots to
            - an extracted entity
            - intent: value pairs
            - a whole message
            or a list of them, where a first match will be picked"""

        return {
            "user_name": [self.from_text()],
            "user_age": [self.from_text()],
            "user_weight": [self.from_text()],
            "user_height": [self.from_text()],
            "user_sex": [self.from_text()],
            "user_scope": [ self.from_text()],
            "user_times_at_gym": [self.from_text()],
            "user_email": [self.from_text()],
        }

    def submit(
            self,
            dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any],
    ) -> List[Dict]:
        """Define what the form has to do
                  after all required slots are filled"""
        user_name = tracker.get_slot('user_name')
        user_age = tracker.get_slot('user_age')
        user_weight = tracker.get_slot('user_weight')
        user_height = tracker.get_slot('user_height')
        user_sex = tracker.get_slot('user_sex')
        user_scope = tracker.get_slot('user_scope')
        user_times_at_gym = tracker.get_slot('user_times_at_gym')
        user_email = tracker.get_slot('user_email')

        mydb = mysql.connector.connect(host="localhost", user="root", passwd="p@ss123",
                                       database="testdatabase")
        query = "insert into users(user_name, user_age, user_weight, user_height, user_sex, user_scope, " \
                "user_times_at_gym, user_email) values ('{}',{},{},{},'{}','{}',{},'{}');".format(user_name, user_age,
                                                                                          user_weight, user_height,
                                                                                          user_sex, user_scope,
                                                                                          user_times_at_gym, user_email)
        print(query)
        cursor = mydb.cursor()

        try:
            result = cursor.execute(query)
            mydb.commit()
            dispatcher.utter_message(template="utter_submit", name=tracker.get_slot('user_name'),
                                     email=tracker.get_slot('user_email'))
        except:
            dispatcher.utter_message("Error!")

        return []

    # def validate(self,dispatcher: CollectingDispatcher,tracker: Tracker,domain: Dict[Text, Any]) -> List[Dict]:
    #     userName=tracker.get_slot('user_name')
    #     dispatcher.utter_message(str(userName))


class ActionRecipeSearch(Action):

    def name(self) -> Text:
        return "action_search_recipe"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

        print("action_search_recipe")

        app_id = 'ff85071b'
        app_key = 'd3efa4ef3ec092ef2cc016bb530897f4'


        recipeIngredients = tracker.get_slot('main_ingredient')

        if not recipeIngredients:
            print(recipeIngredients)
        else:
            print(recipeIngredients)
            url = "https://api.edamam.com/search?q='%s'&app_id=ff85071b&app_key=d3efa4ef3ec092ef2cc016bb530897f4&from=0&to=3" %recipeIngredients
            
            data = json.load(urlopen(url))
            print(url)

            for d in data['hits']:
                recipe = d['recipe']
                recipeName = recipe.get('label')
                ingredients = recipe.get('ingredientLines')
                ingredient = ingredients[0]
                recipeUrl = recipe.get('url')

                dispatcher.utter_message(text="Recipe: {}".format(recipeName))
                dispatcher.utter_message(text="Main ingredients: {}".format(ingredient))
                dispatcher.utter_message(text="More informations here: {}".format(recipeUrl))

                print("Recipe name: ",recipeName)
                print("Recipe ingredients: " , ingredient)


        return []


class ActionMealSearch(Action):

    def name(self) -> Text:
        return "action_search_meal"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:

        print("action_search_meal")
        app_id = '2ae12d8e'
        app_key = '1c41fde6eeefe0f9f39d1100d3eed619'


        #mealIngredients = tracker.latest_message.get("text", "")
        mealIngredients = tracker.get_slot('main_ingredient')

        if not mealIngredients:
            print(mealIngredients)
        else:
            print(mealIngredients)

            editedText = mealIngredients.replace(" ", "%20")
            print(editedText)
            url = "https://api.edamam.com/api/nutrition-data?app_id=2ae12d8e&app_key=1c41fde6eeefe0f9f39d1100d3eed619&ingr=1%20" + editedText
            
            data = json.load(urlopen(url))
            print(url)

            proteins = data.get('calories')
            print(proteins)

            dispatcher.utter_message(text="Your meal is about: {} calories".format(proteins))


        return []


class ActionGymSearch(Action):

    def name(self) -> Text:
        return "action_search_gym"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="Sure! Just a second")



        # Define the API Key.
        API_KEY = 'AIzaSyCa-lfc1U7no3J5vgEquiMrHsWEoQrq0WQ'

        # Define the Client
        gmaps = googlemaps.Client(key = API_KEY)

        # Do a simple nearby search where we specify the location
        # in lat/lon format, along with a radius measured in meters
        places_result  = gmaps.places_nearby(location='44.451893, 26.123038', radius = 2000, open_now =False , type = 'gym')

        time.sleep(2)

        place_result  = gmaps.places_nearby(page_token = places_result['next_page_token'])

        stored_results = []

        counter = 0

        # loop through each of the places in the results, and get the place details.
        for place in places_result['results']:

            if counter == 10:
                break

            # define the place id, needed to get place details. Formatted as a string.
            my_place_id = place['place_id']

            # define the fields you would liked return. Formatted as a list.
            my_fields = ['name','formatted_phone_number','website']

            # make a request for the details.
            places_details  = gmaps.place(place_id= my_place_id , fields= my_fields)

            # print the results of the details, returned as a dictionary.
            print(places_details['result'])

            dispatcher.utter_message(places_details['result'].get("name"))
            dispatcher.utter_message(places_details['result'].get("formatted_phone_number"))
            dispatcher.utter_message(places_details['result'].get("website"))
            dispatcher.utter_message()

            counter = counter + 1

            # store the results in a list object.
            # stored_results.append(places_details['result'])


        return []


class ActionHelloWorldCustom(Action):

    def name(self) -> Text:
        return "action_generate_workout"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        

        dispatcher.utter_message(text="Just a moment")

        print("action_generate_workout")

        mydb = mysql.connector.connect(host="localhost", user="root", passwd="p@ss123",
                                    database="testdatabase")  # # auth_plugin='mysql_native_password'
        userEmail = tracker.get_slot('user_email')
        print(userEmail)
        query = "select user_times_at_gym, user_scope from users where user_email = '{}';".format(userEmail)
        cursor = mydb.cursor()
        print(query)

        no_days_workout=''
        scope_workout = ''

        try:
            x = cursor.execute(query)
            if x == 0:
                print("Sorry, could not find you in the DB")
                dispatcher.utter_message("Sorry I could not find you by your email! :( ")
            else:
                result = cursor.fetchone()
                no_days_workout=result[0]
                scope_workout = result[1]
                print(no_days_workout, scope_workout)
        except:
            dispatcher.utter_message("Sorry I could not find you by your email! :( ")






        noDaysWorkout = no_days_workout
        emailAddress = userEmail
        fileNameWord = 'PersonalisedWorkout.docx'

        document = Document()
        document.add_heading('Workout plan', 0)
        p = document.add_paragraph('Below you will find a detailed version of your personal workout:')

        query = "select exerciseName from workout{} where dayNo = {};".format(noDaysWorkout, 1)
        print(query)
        mydb = mysql.connector.connect(host="localhost", user="root", passwd="p@ss123",
                                    database="testdatabase")  # # auth_plugin='mysql_native_password'

        iterator = 1
        while iterator <= noDaysWorkout:
            query = "select exerciseName from workout{} where dayNo = {};".format(noDaysWorkout, iterator)
            cursor = mydb.cursor()
            print(query)
            x = cursor.execute(query)
            if x == 0:
                print("Sorry, could not find you in the DB")
            else:
                result = cursor.fetchall()
                print(result)

                document.add_heading('Day {}'.format(iterator), level=1)
                print("Creating word document")
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'No Crt'
                hdr_cells[1].text = 'Exercise Name'
                counter = 1
                for exercise in result:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(counter)
                    row_cells[1].text = exercise[0]
                    counter = counter + 1

                print("Creating word document pt2") 
                iterator = iterator + 1

        p = document.add_paragraph('This particular workout split is made in order to first get you used to the volume during '
                                'the first {} weeks. As you can see, you only have {} workout days. Once you get accustomed '
                                'to it, in the final 4 weeks the volume and frequency will go up to 5 workouts a week. '
                                'Really important, whenever you can increase weights from one workout to another, do it, '
                                'as long as you are doing the amount of reps and sets with proper form. Progress is CRUCIAL '
                                ''.format(iterator - 1, iterator - 1))

        p = document.add_paragraph('Before starting the workout do not forget to warm yourself up. You can try the following:   march on the spot: keep going for 3 minutes, heel digs: aim for 60 heel digs in 60 seconds, knee lifts: aim for 30 knee lifts in 30 seconds, shoulder rolls: 2 sets of 10 repetitions, knee bends: 10 repetitions.')

        if scope_workout == 'gain mass':
            p = document.add_paragraph('For GAINING MASS, you will have to workout in the following order: you have to execute the workout above in a slow manner, with big weights and a low number of reps.')
            p = document.add_paragraph('For each exercise you have to do 4 SETS and between 6 to 8 REPS.')
            p = document.add_paragraph('Try to increase the weights as soon as possible. ')
        elif scope_workout == 'keep fit':
            p = document.add_paragraph('For KEEPING FIT you will have to workout in the following order: you have to execute the workout above in a normal manner, with medium weights and an average number of reps.')
            p = document.add_paragraph('For each exercise you have to do 4 SETS and between 8 to 10 REPS.')
            p = document.add_paragraph('Try to do the exericses as correct as possible, described in the videos.')
        elif scope_workout == 'loose weight':
            p = document.add_paragraph('For LOOSING WEIGHT you will have to workout in the following order: you have to execute the workout above in a fast manner, with low weights and a high number of reps.')
            p = document.add_paragraph('For each exercise you have to do 4 SETS and between 14 to 16 REPS.')
            p = document.add_paragraph('Try to do the exericses in an intense manner so that you will burn more calories.')
            p = document.add_paragraph('Before starting the workout to a 30 min set of bicycle run.')


        document.save(fileNameWord)

        email_user = 'thefitnessbot@gmail.com'
        email_password = 'Restanta123'
        email_send = emailAddress

        subject = 'Personal workout'

        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = email_send
        msg['Subject'] = subject

        body = 'Hi there, you will find attached your personalised workout, created by GymChat!'
        msg.attach(MIMEText(body, 'plain'))

        filename = fileNameWord
        attachment = open(filename, 'rb')

        part = MIMEBase('application', 'octet-stream')
        part.set_payload((attachment).read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= " + filename)

        msg.attach(part)
        text = msg.as_string()
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email_user, email_password)


        print("Send mail")
        server.sendmail(email_user, email_send, text)
        server.quit()


        dispatcher.utter_message(text="Please check your mail, you will find the workout there!")

        return []


class ActionHelloWorldCustom(Action):

    def name(self) -> Text:
        return "action_hello_world_program"

    def run(self, dispatcher: CollectingDispatcher,
            tracker: Tracker,
            domain: Dict[Text, Any]) -> List[Dict[Text, Any]]:
        dispatcher.utter_message(text="Sit on the bench holding two dumbbells at shoulder height with an overhand grip. Press the weights up above your head until your arms are fully extended. Return slowly to the start position.")
        dispatcher.utter_message(text="https://www.youtube.com/watch?v=0JfYxMRsUCQ")

        return []
